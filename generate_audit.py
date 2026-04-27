from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

def set_spacing(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)

def add_hr(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_left_border(p, color='B5341C'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)

def add_shading(p, fill='F0EDE8'):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def section_header(doc, number, title):
    add_hr(doc)
    p = doc.add_paragraph()
    set_spacing(p, before=18, after=2)
    r = p.add_run(f"PROBLEME {number}")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xB5, 0x34, 0x1C)
    r.bold = True

    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=10)
    r2 = p2.add_run(title)
    r2.bold = True
    r2.font.size = Pt(15)
    r2.font.color.rgb = RGBColor(0x1A, 0x15, 0x10)

def body_para(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=6)
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x33, 0x2E, 0x28)
    return p

def action_para(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=3, after=3)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    r1 = p.add_run("->  ")
    r1.font.color.rgb = RGBColor(0xB5, 0x34, 0x1C)
    r1.font.size = Pt(11)
    r1.bold = True
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

def label_para(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=10, after=4)
    p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xB5, 0x34, 0x1C)

def quote_para(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=6)
    p.paragraph_format.left_indent = Cm(1.0)
    add_left_border(p, 'B5341C')
    r = p.add_run('"' + text + '"')
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x50, 0x48)

def code_para(doc, text, good=False):
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=4)
    p.paragraph_format.left_indent = Cm(0.5)
    add_shading(p, 'F0EDE8')
    color = RGBColor(0x22, 0x88, 0x44) if good else RGBColor(0xB5, 0x34, 0x1C)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)
    r.font.color.rgb = color

# ── TITLE PAGE ──────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p, before=40, after=4)
r = p.add_run("AUDIT DU SITE")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1A, 0x15, 0x10)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p2, before=0, after=4)
r2 = p2.add_run("antoninatger.com")
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0xB5, 0x34, 0x1C)
r2.italic = True

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p3, before=0, after=30)
r3 = p3.add_run("Ce qui empeche les visiteurs de commander une intervention")
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor(0x66, 0x60, 0x58)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(p4, before=0, after=0)
r4 = p4.add_run("Avril 2026")
r4.font.size = Pt(10)
r4.font.color.rgb = RGBColor(0xAA, 0xA5, 0x9C)

doc.add_page_break()

# ── SOMMAIRE ─────────────────────────────────────────────
p = doc.add_paragraph()
set_spacing(p, before=0, after=16)
r = p.add_run("SOMMAIRE")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0xB5, 0x34, 0x1C)

sommaire = [
    ("1.", "Probleme structurel : la page d'accueil parle du mauvais sujet"),
    ("2.", "La biographie parle de vous, pas du client"),
    ("3.", "Un seul temoignage pour 122 structures"),
    ("4.", "Vos meilleurs credentials ne sont pas sur la page Interventions"),
    ("5.", "La formation entreprise est sous-developpee"),
    ("6.", "Le contact est flou et distant"),
    ("7.", "Le Fakemetre est une coquille vide"),
    ("8.", "Zone geographique inconnue"),
    ("9.", "La page Collaborations dit journaliste mais pas intervenant"),
    ("--", "Tableau recapitulatif"),
]

for num, titre in sommaire:
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=3)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{num}  ")
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0xB5, 0x34, 0x1C)
    r1.bold = True
    r2 = p.add_run(titre)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x33, 0x2E, 0x28)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# PROBLEME 1
# ═══════════════════════════════════════════════════════
section_header(doc, "1 -- Structurel", "La page d'accueil parle du mauvais sujet")
body_para(doc, "La page principale s'intitule 'Publications' et s'ouvre sur le roman Interfeel. Un directeur d'ecole, un responsable de mediatheque ou un DRH qui atterrit la ne comprend pas en trois secondes que vous pouvez intervenir chez lui.")
body_para(doc, "Votre vitrine actuelle vend un auteur, pas un intervenant. L'ordre des priorites est inverse : le livre passe avant le service que vous cherchez a vendre.")
label_para(doc, "Constat")
action_para(doc, "La section Interventions de la page Publications est secondaire et courte.")
action_para(doc, "Il n'y a pas de hierarchie visuelle qui oriente un visiteur venu pour 'reserver une conference'.")
action_para(doc, "Le titre de la page (Publications) n'invite pas un acheteur de services a rester.")
label_para(doc, "A corriger")
action_para(doc, "Faire d'une page Interventions la page centrale ou la page d'accueil.")
action_para(doc, "Deplacer les romans en page secondaire, accessibles depuis la navigation.")
action_para(doc, "Ajouter un titre de type : 'Conferences, ateliers et formations -- Fake news, esprit critique, ecriture'.")

# ═══════════════════════════════════════════════════════
# PROBLEME 2
# ═══════════════════════════════════════════════════════
section_header(doc, "2", "La biographie parle de vous, pas du client")
body_para(doc, "La biographie actuelle commence par :")
quote_para(doc, "Je suis tout d'abord ecrivain de science-fiction.")
body_para(doc, "C'est la premiere phrase que lit quelqu'un qui cherche a reserver une intervention. Elle n'a aucun pouvoir de conviction. Une biographie orientee 'vente de service' s'ouvre sur le probleme du client -- puis propose une reponse.")
label_para(doc, "Exemple de reformulation d'accroche (a adapter)")
p = doc.add_paragraph()
set_spacing(p, before=4, after=8)
p.paragraph_format.left_indent = Cm(1.0)
add_left_border(p, 'CCCCCC')
r = p.add_run("Fake news, complotisme, arnaques en ligne : comment enseigner l'esprit critique a des eleves qui ne font plus confiance a rien ? C'est la question qui a oriente mon parcours -- d'abord auteur de science-fiction, puis chercheur en psychologie, aujourd'hui intervenant dans plus de 120 structures en France.")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x44, 0x40, 0x38)
label_para(doc, "A corriger")
action_para(doc, "Commencer par le probleme du lecteur, pas par votre identite d'auteur.")
action_para(doc, "Mentionner les chiffres cles des les premieres lignes (86 etablissements, 36 mediathequess, TEDx, Assemblee Nationale).")
action_para(doc, "Terminer par un appel a l'action, pas par une liste de cours suivis.")

# ═══════════════════════════════════════════════════════
# PROBLEME 3
# ═══════════════════════════════════════════════════════
section_header(doc, "3", "Un seul temoignage pour 122 structures")
body_para(doc, "Sur la page Interventions, il n'y a qu'un seul temoignage (Samuel B., Lycee Charles de Gaulle). Pourtant, l'accordeon 'Ils m'ont fait confiance' affiche 86 etablissements scolaires et 36 mediathequess. C'est la preuve sociale la plus puissante du site -- et elle est cachee derriere un clic.")
label_para(doc, "Ce que chaque profil veut entendre")
action_para(doc, "Bibliothecaire / responsable mediatheque : 'L'intervention etait facile a organiser, le public a bien reagi, on a eu des retours tres positifs.'")
action_para(doc, "CPE / proviseur : 'Les eleves sont restes attentifs, il a su capter meme les classes difficiles.'")
action_para(doc, "DRH / responsable formation : 'Nos equipes ont reparte avec des reflexes concrets. Le format s'est adapte a notre contexte.'")
label_para(doc, "A corriger")
action_para(doc, "Collecter 5 a 8 temoignages de profils varies (enseignant, bibliothecaire, CPE, RH).")
action_para(doc, "Les afficher directement, sans accordeon, en haut de page.")
action_para(doc, "Inclure le nom complet, le role et l'etablissement (avec accord) pour credibiliser.")

# ═══════════════════════════════════════════════════════
# PROBLEME 4
# ═══════════════════════════════════════════════════════
section_header(doc, "4", "Vos meilleurs credentials ne sont pas sur la page Interventions")
body_para(doc, "Vos quatre preuves de legitimite les plus fortes sont eparpillees sur des pages secondaires que la plupart des visiteurs ne consulteront pas :")
label_para(doc, "Credentials enteres")
action_para(doc, "Auditionne par l'Assemblee Nationale --> dans 'Couverture mediatique'")
action_para(doc, "TEDx --> dans 'Collaborations'")
action_para(doc, "Prix europeen Voices 2025 (Zagreb) --> dans 'Collaborations'")
action_para(doc, "France 3 Auvergne-Rhone-Alpes --> dans 'Couverture mediatique'")
label_para(doc, "A corriger")
action_para(doc, "Creer une bande de credibilite en haut de page Interventions : TEDx - Assemblee Nationale - Prix Voices 2025 - France 3 - 122 structures.")
action_para(doc, "Ajouter un encart 'Vu sur' ou 'References' avec logos ou noms visibles sans clic.")

# ═══════════════════════════════════════════════════════
# PROBLEME 5
# ═══════════════════════════════════════════════════════
section_header(doc, "5", "La formation entreprise est sous-developpee")
body_para(doc, "Le format propose est de 30 a 45 minutes. C'est tres court pour une formation professionnelle. Il n'y a aucun temoignage d'un responsable RH ou d'un manager. Pas d'exemple de secteur, de taille d'equipe, de retour mesurable.")
label_para(doc, "Ce qui manque cote entreprise")
action_para(doc, "Un ou deux formats supplementaires : demi-journee (3h), journee complete.")
action_para(doc, "Un cas concret ou une etude de cas : secteur, probleme rencontre, resultat.")
action_para(doc, "Un temoignage d'un responsable de formation ou d'un directeur.")
action_para(doc, "Une mention de la prise en charge OPCO si applicable.")
action_para(doc, "Les secteurs qui vous ont deja sollicite (presse, tech, collectivites, education...).")

# ═══════════════════════════════════════════════════════
# PROBLEME 6
# ═══════════════════════════════════════════════════════
section_header(doc, "6", "Le contact est flou et distant")
body_para(doc, "Tous les boutons CTA pointent vers antoninatger.com/contact/ (site WordPress externe). Il n'y a pas d'email, de numero de telephone, ni de formulaire integre dans les pages auditees.")
body_para(doc, "La seule mention de tarif sur tout le site est : 'Pour toute demande de tarif, vous pouvez utiliser le formulaire de contact disponible sur ce site.' C'est insuffisant pour declencher une prise de contact.")
label_para(doc, "A corriger")
action_para(doc, "Afficher une adresse email directement visible sur la page Interventions.")
action_para(doc, "Indiquer un delai de reponse (ex : 'Je reponds sous 48h').")
action_para(doc, "Envisager une tarification indicative ou une plage de prix : 'A partir de X euros' rassure sans engager.")
action_para(doc, "Ajouter un encart contact en bas de chaque fiche d'intervention.")

# ═══════════════════════════════════════════════════════
# PROBLEME 8
# ═══════════════════════════════════════════════════════
section_header(doc, "7", "Le Fakemetre est une coquille vide")
body_para(doc, "La section Fakemetre (bas de index.html) a un titre, un sous-titre, une description -- mais aucun lien, aucun bouton, aucune demo accessible. Elle ressemble a du contenu inacheve mis en ligne.")
body_para(doc, "Or c'est exactement le type d'outil qui peut convaincre un enseignant ou un bibliothecaire : il peut l'essayer, le montrer a ses eleves, et realiser la valeur de votre approche avant meme de vous contacter.")
label_para(doc, "A corriger")
action_para(doc, "Si l'outil est pret : ajouter un bouton 'Tester le Fakemetre' ou l'integrer directement.")
action_para(doc, "Si l'outil n'est pas pret : supprimer la section ou la remplacer par 'En cours de developpement'.")

# ═══════════════════════════════════════════════════════
# PROBLEME 9
# ═══════════════════════════════════════════════════════
section_header(doc, "8", "Zone geographique inconnue")
body_para(doc, "Nulle part sur le site il n'est indique si vous intervenez dans toute la France, principalement en region lyonnaise, ou a l'international. Pour un proviseur a Bordeaux ou un bibliothecaire a Lille, c'est une question bloquante.")
label_para(doc, "A corriger")
action_para(doc, "Ajouter une ligne dans 'Informations pratiques' : ex. 'J'interviens dans toute la France, principalement en Auvergne-Rhone-Alpes. Des deplacements a l'etranger sont possibles (cf. Zagreb 2025).'")

# ═══════════════════════════════════════════════════════
# PROBLEME 10
# ═══════════════════════════════════════════════════════
section_header(doc, "9", "La page Collaborations dit journaliste mais pas intervenant")
body_para(doc, "La page collaborations.html est excellente comme portfolio mediatique. Mais elle est encadree comme une liste de medias et de podcasts, pas comme une preuve que vous intervenez deja dans des structures similaires a celles du visiteur.")
body_para(doc, "Les interventions a la mediatheque Mejanes, au Departement du Val d'Oise, a l'HEP Lyon -- ce sont exactement les preuves qu'un acheteur cherche. Mais le cadrage 'collaborations medias' les noie dans le flux.")
label_para(doc, "A corriger")
action_para(doc, "Separer clairement : 'Conferences & interventions publiques' vs 'Medias & podcasts'.")
action_para(doc, "Les conferences publiques enregistrees (Mejanes, Val d'Oise, HEP Lyon) devraient etre sur la page Interventions, pas seulement ici.")

# ═══════════════════════════════════════════════════════
# TABLEAU RECAPITULATIF
# ═══════════════════════════════════════════════════════
add_hr(doc)
p = doc.add_paragraph()
set_spacing(p, before=18, after=12)
r = p.add_run("TABLEAU RECAPITULATIF")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x1A, 0x15, 0x10)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr = table.rows[0].cells
hdr[0].text = "Probleme"
hdr[1].text = "Consequence"
hdr[2].text = "Priorite"

for cell in hdr:
    for para in cell.paragraphs:
        para.runs[0].bold = True
        para.runs[0].font.size = Pt(10)
        para.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A1510')
    cell._tc.get_or_add_tcPr().append(shd)

rows_data = [
    ("Page d'accueil orientee auteur, pas intervenant", "Visiteurs qui repartent sans contacter", "URGENT"),
    ("Biographie centree sur vous, pas le client", "Pas de connexion avec l'acheteur", "URGENT"),
    ("Un seul temoignage pour 122 structures", "Manque de preuve sociale", "URGENT"),
    ("TEDx / Assemblee Nationale / Voices 2025 enteres", "Credibilite invisible", "URGENT"),
    ("Formation entreprise sous-developpee", "Perd des clients professionnels", "IMPORTANT"),
    ("Contact flou, pas d'email direct", "Friction avant prise de contact", "IMPORTANT"),
    ("Fakemetre vide", "Outil promis, non livre", "IMPORTANT"),
    ("Zone geographique inconnue", "Doute pour les visiteurs hors Lyon", "MOYEN"),
    ("Page Collaborations mal structuree", "Preuves de terrain noyees", "MOYEN"),
]

for probleme, impact, priorite in rows_data:
    row = table.add_row().cells
    row[0].text = probleme
    row[1].text = impact
    row[2].text = priorite
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x22, 0x1E, 0x18)

p = doc.add_paragraph()
set_spacing(p, before=20, after=0)
r = p.add_run("Document produit le 27 avril 2026  |  antoninatger.com")
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = RGBColor(0xAA, 0xA5, 0x9C)

path = r"C:\Users\ASUS\Documents\Web apps\Wordpress\Bibliographie\audit-site-antonin-atger.docx"
doc.save(path)
print(f"Document cree : {path}")
