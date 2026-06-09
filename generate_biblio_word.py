from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Marges ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

ROUGE  = RGBColor(0xB5, 0x34, 0x1C)
NOIR   = RGBColor(0x11, 0x11, 0x11)
GRIS   = RGBColor(0x66, 0x66, 0x66)

def set_font(run, size, bold=False, italic=False, color=None):
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_eyebrow(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    set_font(run, 8, bold=True, color=ROUGE)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_h1(doc, text):
    p = doc.add_heading(level=1)
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)
    run = p.add_run(text)
    set_font(run, 26, bold=True, color=NOIR)
    return p

def add_section_break(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(f"── {text.upper()} ──")
    set_font(run, 8, bold=True, color=RGBColor(0x99, 0x99, 0x99))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h2(doc, text):
    p = doc.add_heading(level=2)
    p.clear()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 14, bold=True, italic=True, color=NOIR)
    return p

def add_body(doc, text, space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, 11, color=color or RGBColor(0x33, 0x33, 0x33))
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'DDDDDD')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_timeline_item(doc, year, title, type_tag, text, italic_title=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(0)
    r_year = p.add_run(f"{year}  ")
    set_font(r_year, 10, italic=True, color=RGBColor(0x88, 0x88, 0x88))
    r_title = p.add_run(title)
    set_font(r_title, 12, bold=True, italic=italic_title, color=NOIR)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after  = Pt(2)
    p2.paragraph_format.left_indent  = Cm(1.5)
    r_tag = p2.add_run(type_tag)
    set_font(r_tag, 8, bold=True, color=RGBColor(0x88, 0x88, 0x88))

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(1)
    p3.paragraph_format.space_after  = Pt(4)
    p3.paragraph_format.left_indent  = Cm(1.5)
    r_text = p3.add_run(text)
    set_font(r_text, 10.5, color=RGBColor(0x44, 0x44, 0x44))

# ══════════════════════════════════════════
#  EN-TÊTE
# ══════════════════════════════════════════
add_eyebrow(doc, "Antonin Atger")
add_h1(doc, "Bibliographie & Publications")
add_hr(doc)

# ══════════════════════════════════════════
#  BIOGRAPHIE
# ══════════════════════════════════════════
add_section_break(doc, "Biographie")

add_body(doc,
    "Écrivain de science-fiction, Antonin Atger est l'auteur de la série Interfeel, dans laquelle il "
    "imagine un réseau social permettant de partager ses émotions. Titulaire d'un premier master en "
    "sciences politiques, il a effectué un second master en psychologie, dont la thèse porte sur les "
    "liens entre psychologie, complotisme et stratégie politique.",
    space_after=8)

add_body(doc,
    "Ces recherches l'ont conduit à se spécialiser dans la désinformation en ligne et les réseaux sociaux — "
    "connaissances qu'il vulgarise auprès des jeunes lors de ses interventions en milieu scolaire et professionnel.",
    space_after=8)

# ══════════════════════════════════════════
#  SÉRIE INTERFEEL
# ══════════════════════════════════════════
add_section_break(doc, "Romans publiés")

# Tome 1
add_h2(doc, "Interfeel — Tome 1")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Éditions Pocket Jeunesse  ·  Roman d'anticipation")
set_font(r, 9, color=RGBColor(0x88, 0x88, 0x88))
add_body(doc,
    "Nathan et ses amis sont en permanence connectés à Interfeel, un réseau social qui permet de partager "
    "ses émotions. Mais un événement tragique va bouleverser Nathan. Fasciné par Élizabeth, une "
    "« sans-Réseau » qui vit en marge de la société, il voit toutes ses certitudes vaciller. Ce que les "
    "deux adolescents découvriront pourrait bien changer le monde à jamais…",
    space_after=4)
add_body(doc,
    "Né d'un concours Welovewords remporté en 2012. Lauréat du CLA Books Awards 2019 et du Prix Chimère 2019.",
    color=RGBColor(0x77, 0x77, 0x77), space_after=10)

add_hr(doc)

# Tome 2
add_h2(doc, "Interfeel — Tome 2 : Les Résistants")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Éditions Pocket Jeunesse  ·  Roman d'anticipation")
set_font(r, 9, color=RGBColor(0x88, 0x88, 0x88))
add_body(doc,
    "Après la rupture d'Interfeel, l'équipe « Résistance » est dispersée. Adila et Nadem se voient confier "
    "une mission : traverser le monde pour retrouver le créateur d'Interfeel, seule personne capable de "
    "stopper la machination en cours. Le Tatoueur tisse sa toile jour après jour — de plus en plus puissant, "
    "incontrôlable, et dangereux. Une seule personne peut l'arrêter : Nathan Ethanin. Mais depuis six mois, "
    "il a disparu…",
    space_after=10)

add_hr(doc)

# Tome 3
add_h2(doc, "Interfeel — Tome 3 : L'Odyssée")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Éditions Pocket Jeunesse  ·  Roman d'anticipation")
set_font(r, 9, color=RGBColor(0x88, 0x88, 0x88))
add_body(doc,
    "Lorsque Adila et Nadir arrivent aux portes du pays rouge, le plus dur reste à venir. Hanek et le "
    "professeur ont retrouvé l'Organique, membre légendaire de l'organisation F.A.N.T.O.M.E. Pour retrouver "
    "Nathan, les ennemis d'hier doivent s'allier contre la menace grandissante et invisible du Tatoueur. "
    "Le sort de tous est désormais entre les mains de quelques-uns…",
    space_after=10)

# ══════════════════════════════════════════
#  RÉCOMPENSES
# ══════════════════════════════════════════
add_section_break(doc, "Récompenses")

for prix in ["CLA Books Awards 2019", "Prix Chimère 2019"]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(prix)
    set_font(r, 11, bold=True, color=NOIR)

# ══════════════════════════════════════════
#  PRESSE
# ══════════════════════════════════════════
add_section_break(doc, "Revue de presse")

presse = ["Le Parisien", "Var Matin", "Le Progrès", "Biblioteca Magazine", "Mon Quotidien"]
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Couverture médiatique : " + " · ".join(presse))
set_font(r, 11, color=RGBColor(0x33, 0x33, 0x33))

# ══════════════════════════════════════════
#  PARCOURS LITTÉRAIRE
# ══════════════════════════════════════════
add_section_break(doc, "Parcours littéraire")

add_timeline_item(doc, 2012, "Sélection ACLA",
    "Nouvelle · Association Culture Loisir Antibes",
    "Sélectionné pour publication dans le recueil de l'Association Culture Loisir Antibes sur un sujet imposé particulièrement exigeant.")

add_timeline_item(doc, 2012, "L'Indifférence", "Nouvelle · Sorbonne, Master Édition",
    "Sélectionné parmi les candidatures à cet appel à textes sur le thème du fait divers. "
    "Présence au Salon du Livre de Paris.", italic_title=True)

add_timeline_item(doc, 2009, "Colloque — Le Rythme, le Mot et l'Image",
    "Communication académique · University of Victoria, Vancouver",
    "Communication sur les corrélations entre réalité, livre et cinéma à travers Entre les Murs "
    "(François Bégaudeau). Un échange avec le réalisateur Laurent Cantet a pu être établi dans ce cadre.", italic_title=True)

add_timeline_item(doc, 2009, "Les Étoiles — publication en recueil",
    "Recueil · Quai du Polar",
    "La nouvelle primée en 2008 est republiée dans le recueil rassemblant les lauréats des quatre dernières éditions du festival.", italic_title=True)

add_timeline_item(doc, 2008, "Meyzieu, quand la ville danse",
    "Ouvrage collectif",
    "Collaboration avec le photographe Jean-François Marin : plus d'une centaine d'entretiens menés auprès "
    "des habitants, tissés en une trame narrative. Vendu à plus de 2 000 exemplaires.", italic_title=True)

add_timeline_item(doc, 2008, "Lauréat du Quai du Polar",
    "Concours · Nouvelle",
    "Première participation à un concours littéraire, premier prix. Thème imposé : « Rendez-vous sur les quais ». "
    "Nouvelle intitulée Les Étoiles.")

# ══════════════════════════════════════════
#  INTERVENTIONS (résumé)
# ══════════════════════════════════════════
add_section_break(doc, "Interventions")

interventions = [
    ("Fake News et esprit critique",
     "Identifier les fausses informations, comprendre leurs mécanismes de diffusion et développer les réflexes de vérification."),
    ("Désinformation, au-delà de la Fake News",
     "Décontextualisation, biais cognitifs, propagande algorithmique et stratégies politiques."),
    ("Rhétorique et esprit critique",
     "Déconstruire un discours, identifier les sophismes et affûter sa pensée critique."),
    ("Ateliers d'écriture",
     "Ateliers créatifs autour de la nouvelle et de la science-fiction."),
    ("Le piège des influenceurs",
     "Mécanismes d'influence sur les réseaux sociaux et vecteurs de désinformation."),
    ("La magie des histoires",
     "Le storytelling au service de la désinformation, développer un regard critique."),
    ("Formation contre l'arnaque en ligne",
     "Phishing, faux profils, escroqueries aux sentiments, fraudes à la livraison."),
]

for title, desc in interventions:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(title)
    set_font(r, 11, bold=True, color=NOIR)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(3)
    p2.paragraph_format.left_indent  = Cm(0.5)
    r2 = p2.add_run(desc)
    set_font(r2, 10, color=RGBColor(0x55, 0x55, 0x55))

# ── Sauvegarde ──
out = r"C:\Users\ASUS\Documents\Web apps\Wordpress\Bibliographie\Bibliographie_Antonin_Atger.docx"
doc.save(out)
print(f"Document généré : {out}")
